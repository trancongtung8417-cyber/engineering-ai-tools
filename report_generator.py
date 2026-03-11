import streamlit as st
from openai import OpenAI
from docx import Document
from io import BytesIO

st.set_page_config(page_title="AI Report Generator", page_icon="📝")

st.title("📝 AI Technical Report Generator")
st.write("Convert site observations into professional Word reports.")

with st.sidebar:
    api_key = st.text_input("Enter OpenAI API Key", type="password")

# Form nhập liệu
field_notes = st.text_area("Enter raw field notes or transcription:", placeholder="e.g., motor overheating, 85C, noise from bearings, need replacement by Friday.")
project_name = st.text_input("Project Name/ID", "SITE-INSPECTION-001")

if st.button("Generate Professional Report"):
    if not api_key:
        st.error("Missing API Key!")
    elif not field_notes:
        st.warning("Please enter some notes first.")
    else:
        with st.spinner("AI is drafting your report..."):
            client = OpenAI(api_key=api_key)
            
            # Gọi AI để "chuốt" văn bản
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a Professional Technical Consultant. Convert rough notes into a formal report with sections: Subject, Status, Observations, Root Cause, and Recommendations. Use passive voice and professional engineering English."},
                    {"role": "user", "content": field_notes}
                ]
            )
            report_content = response.choices[0].message.content
            
            # Hiển thị kết quả ra màn hình
            st.subheader("Preview")
            st.markdown(report_content)
            
            # Tạo file Word để tải về
            doc = Document()
            doc.add_heading(f"TECHNICAL REPORT: {project_name}", 0)
            doc.add_paragraph(report_content)
            
            bio = BytesIO()
            doc.save(bio)
            
            st.download_button(
                label="📥 Download Report (.docx)",
                data=bio.getvalue(),
                file_name=f"{project_name}_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )